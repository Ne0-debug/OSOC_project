from fastapi import FastAPI
from fastapi.responses import FileResponse
from pydantic import BaseModel
from pathlib import Path
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
import os
import re
import json
import uuid

load_dotenv()

app = FastAPI()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model = genai.GenerativeModel("gemini-3.6-flash")

# --- Load all chunk files at startup ---
CHUNKS_DIR = Path(__file__).parent / "docs_temp" / "chunks"

def load_chunks():
    chunks = []
    for file_path in CHUNKS_DIR.glob("*.md"):
        text = file_path.read_text(encoding="utf-8")
        first_line = text.splitlines()[0]
        topic = first_line.replace("# Topic:", "").strip()
        chunks.append({
            "filename": file_path.name,
            "topic": topic,
            "text": text
        })
    return chunks

CHUNKS = load_chunks()


@app.get("/test")
def test():
    response = model.generate_content("In one sentence, what is the RTI Act, 2005?")
    return {"answer": response.text}


STOPWORDS = {"the", "is", "of", "a", "an", "to", "in", "for", "and", "or", "what", "how", "do", "does", "i", "my", "can", "you", "me", "on", "at", "with"}

def score_chunk(query: str, chunk: dict) -> int:
    query_words = set(re.findall(r"\w+", query.lower())) - STOPWORDS
    chunk_words = set(re.findall(r"\w+", (chunk["topic"] + " " + chunk["text"]).lower())) - STOPWORDS
    return len(query_words & chunk_words)

def select_relevant_chunks(query: str, top_n: int = 3):
    scored = [(score_chunk(query, chunk), chunk) for chunk in CHUNKS]
    scored.sort(key=lambda x: x[0], reverse=True)
    relevant = [chunk for score, chunk in scored if score >= 2]
    return relevant[:top_n]


class AskRequest(BaseModel):
    query: str

@app.post("/ask")
def ask(request: AskRequest):
    relevant_chunks = select_relevant_chunks(request.query)

    if not relevant_chunks:
        return {"answer": "I don't have information on this in my current sources.", "sources": []}

    context = "\n\n".join(
        f"[Source: {c['filename']}]\n{c['text']}" for c in relevant_chunks
    )

    prompt = f"""You are a helpful legal assistant. Answer the user's question using ONLY the context below. 
Cite which source file(s) you used at the end of your answer, like "(Source: filename.md)".
If the context does not contain the answer, say "I don't have information on this" — do not guess.

Context:
{context}

Question: {request.query}

Answer:"""

    response = model.generate_content(prompt)
    return {
        "answer": response.text,
        "sources": [c["filename"] for c in relevant_chunks]
    }


# --- Document generation ---

RTI_FIELDS = ["applicant_name", "applicant_address", "department_name", "subject", "information_requested", "date"]
TENANCY_FIELDS = ["applicant_name", "applicant_address", "landlord_name", "property_address", "deposit_amount", "date_vacated", "specific_ask"]

class GenerateDocRequest(BaseModel):
    conversation: str
    document_type: str  # "rti" or "tenancy_complaint"

def extract_fields(conversation: str, document_type: str):
    fields = RTI_FIELDS if document_type == "rti" else TENANCY_FIELDS

    prompt = f"""Extract the following fields from this conversation. Return ONLY valid JSON, no markdown, no explanation.
If a field is missing from the conversation, set its value to null.

Fields needed: {fields}

Conversation:
{conversation}

Return format: {{"field_name": "value", ...}}"""

    response = model.generate_content(prompt)
    text = response.text.strip()
    text = text.replace("```json", "").replace("```", "").strip()
    return json.loads(text)


def build_rti_docx(fields: dict) -> str:
    doc = Document()
    doc.add_heading("Application under Right to Information Act, 2005", level=1)
    doc.add_paragraph(f"To,\nThe Public Information Officer,\n{fields.get('department_name') or '[Department Name]'}")
    doc.add_paragraph(f"From,\n{fields.get('applicant_name') or '[Applicant Name]'}\n{fields.get('applicant_address') or '[Applicant Address]'}")
    doc.add_paragraph(f"Date: {fields.get('date') or '[Date]'}")
    doc.add_heading("Subject:", level=2)
    doc.add_paragraph(fields.get('subject') or '[Subject not provided]')
    doc.add_heading("Information Requested:", level=2)
    doc.add_paragraph(fields.get('information_requested') or '[Details not provided]')
    doc.add_paragraph("\nI request that the above information be provided to me under the Right to Information Act, 2005.")
    doc.add_paragraph(f"\nSincerely,\n{fields.get('applicant_name') or '[Applicant Name]'}")

    filename = f"rti_application_{uuid.uuid4().hex[:8]}.docx"
    filepath = f"generated_docs/{filename}"
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
    return filepath


def build_tenancy_docx(fields: dict) -> str:
    doc = Document()
    doc.add_heading("Complaint Regarding Security Deposit", level=1)
    doc.add_paragraph("To,\nThe Rent Authority")
    doc.add_paragraph(f"From,\n{fields.get('applicant_name') or '[Applicant Name]'}\n{fields.get('applicant_address') or '[Applicant Address]'}")
    doc.add_heading("Details:", level=2)
    doc.add_paragraph(f"Landlord Name: {fields.get('landlord_name') or '[Not provided]'}")
    doc.add_paragraph(f"Property Address: {fields.get('property_address') or '[Not provided]'}")
    doc.add_paragraph(f"Deposit Amount: {fields.get('deposit_amount') or '[Not provided]'}")
    doc.add_paragraph(f"Date Vacated: {fields.get('date_vacated') or '[Not provided]'}")
    doc.add_heading("Request:", level=2)
    doc.add_paragraph(fields.get('specific_ask') or '[Not provided]')
    doc.add_paragraph(f"\nSincerely,\n{fields.get('applicant_name') or '[Applicant Name]'}")

    filename = f"tenancy_complaint_{uuid.uuid4().hex[:8]}.docx"
    filepath = f"generated_docs/{filename}"
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
    return filepath


@app.post("/generate-document")
def generate_document(request: GenerateDocRequest):
    fields = extract_fields(request.conversation, request.document_type)

    missing = [f for f in fields if fields[f] is None]
    if missing:
        return {"needs_more_info": True, "missing_fields": missing, "extracted_so_far": fields}

    if request.document_type == "rti":
        filepath = build_rti_docx(fields)
    elif request.document_type == "tenancy_complaint":
        filepath = build_tenancy_docx(fields)
    else:
        return {"error": "document_type must be 'rti' or 'tenancy_complaint'"}

    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(filepath)
    )